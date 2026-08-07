from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION

OUT = 'FPGA_Drone_Accelerator_Product_Scope_Revised.docx'
BLUE = '2E74B5'
NAVY = '1F4D78'
LIGHT = 'F2F4F7'
INK = '202020'

def set_font(run, size=11, bold=None, color=INK, italic=None):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def borders(table):
    tblPr = table._tbl.tblPr
    tb = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:color'),'D9E1F2'); tb.append(e)
    tblPr.append(tb)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None: node = OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'),'dxa')

def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW')
    if tblW is None: tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths))); tblW.set(qn('w:type'),'dxa')
    ind = OxmlElement('w:tblInd'); ind.set(qn('w:w'),'120'); ind.set(qn('w:type'),'dxa'); tblPr.append(ind)
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths): col.set(qn('w:w'),str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.first_child_found_in('w:tcW')
            if tcW is None: tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
            tcW.set(qn('w:w'),str(width)); tcW.set(qn('w:type'),'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    borders(table)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); set_font(r)
    return p

def add_para(doc, text='', bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead); set_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):]); set_font(r)
    else:
        r = p.add_run(text); set_font(r)
    return p

def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    r = p.add_run(text); set_font(r, {1:16,2:13,3:12}[level], bold=True, color=BLUE if level < 3 else NAVY)
    return p

def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]; shade(cell, LIGHT)
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(label), 10, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(value), 10)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table

doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)

normal = doc.styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); normal.font.size=Pt(11)
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.10
for level, size, before, after in [(1,16,16,8),(2,13,12,6),(3,12,8,4)]:
    st=doc.styles[f'Heading {level}']; st.font.name='Calibri'; st._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(BLUE if level<3 else NAVY); st.font.bold=True
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after)

footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(footer.add_run('CoreLink | Product scope revision | July 2026'), 8, color='6B7280')

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4)
set_font(p.add_run('PRODUCT DECISION MEMO'), 10, bold=True, color=BLUE)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
set_font(p.add_run('FPGA Co-Processor for Adaptive Utility-Inspection Drones'), 23, bold=True, color=NAVY)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(16)
set_font(p.add_run('Revised MVP, thesis contribution and market-entry plan'), 13, color='4B5563')
for label, value in [('Decision', 'Proceed with a power-line inspection co-processor, not a generic defect-detection card.'), ('Scope', 'One camera stream + MAVLink/PX4 companion interface; advisory outputs only in MVP.'), ('Thesis claim', 'Quantized FPGA perception and capture-quality logic with bounded sensor-to-decision latency, evaluated against Jetson and an NPU baseline.'), ('Date', '26 July 2026')]:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    set_font(p.add_run(label+': '), 10, bold=True, color=NAVY); set_font(p.add_run(value), 10)

heading(doc, '1. Recommendation', 1)
add_para(doc, 'Keep energy-infrastructure inspection as the vertical, but change the product from a “defect detector on an FPGA” to an adaptive inspection co-processor for overhead power lines. The MVP detects inspection targets (for example, insulators, clamps and conductors), evaluates whether each frame is usable, and emits deterministic capture / re-capture / slow-down recommendations to the mission stack. It never directly commands flight controls in the MVP.')
add_para(doc, 'This is a materially stronger thesis and startup wedge. A fixed CNN detector alone is easy to compare against low-power NPUs and does not prove why an FPGA is needed. The proposed product uses the FPGA where it is credible: custom camera-to-logic dataflow, bounded latency, pipeline composition, and the ability to evolve sensor and model interfaces without a silicon respin.')

heading(doc, '2. The MVP: adaptive capture for power-line inspection', 1)
add_para(doc, 'The buyer’s practical problem is not merely finding a defect after landing. Inspection teams lose time when imagery is blurred, badly framed, duplicated, or missing the component required for a maintenance decision. The MVP makes the aircraft an active data-collection tool: it decides in flight whether a component has been seen at sufficient quality and produces an auditable record of candidate defects and coverage.')
add_table(doc, ['MVP function', 'What it does', 'Why it matters'], [
    ('Target localization', 'Detect conductor, insulator and hardware regions in RGB video.', 'Anchors capture and supports component-level coverage.'),
    ('Quality / coverage scoring', 'Score focus, exposure, framing, scale and repeat coverage against a mission template.', 'Reduces unusable imagery and unnecessary repeat flights.'),
    ('Candidate anomaly flag', 'Run a lightweight classifier on the selected crop; mark, do not diagnose, suspected damage.', 'Provides immediate review priority without overclaiming accuracy.'),
    ('Deterministic advisory output', 'Publish timestamped MAVLink / ROS 2 messages: capture, retry, slow-down, candidate.', 'Integrates with existing PX4/ArduPilot workflows while keeping the pilot or mission controller in charge.')
], [2200, 3900, 3260])

heading(doc, '3. Product boundary: what is deliberately out of scope', 1)
add_bullet(doc, 'No autonomous BVLOS claim, safety certification claim, or direct flight-control actuation in the MVP. Those add regulatory and liability burden before the value proposition is proven.')
add_bullet(doc, 'No multi-sensor SLAM, LiDAR pre-processing or video encoder in v1. They are credible future FPGA workloads, but combining them now turns the thesis into an integration programme rather than a publishable experiment.')
add_bullet(doc, 'No promise that the FPGA wins TOPS/W against an NPU for the same fixed detector. The evaluation must state where it loses as well as where deterministic pipeline behavior and interface flexibility win.')

heading(doc, '4. Thesis contribution and publishable research question', 1)
add_para(doc, 'Research question: Can a co-designed FPGA vision pipeline improve the timeliness and usable-data yield of close-range utility inspection, while retaining bounded end-to-end latency and competitive energy per accepted inspection observation?')
add_table(doc, ['Contribution', 'Evidence required'], [
    ('Co-designed workload', 'INT8 target detector plus image-quality / coverage logic mapped as a streaming FPGA pipeline.'),
    ('Real-time systems result', 'Measure latency distribution and worst-case bound from camera ingress to advisory message under sustained load.'),
    ('Operational metric', 'Report accepted target observations per flight minute, re-capture rate, and candidate-review precision; not only mAP and FPS.'),
    ('Fair accelerator comparison', 'Compare FPGA, Jetson and Hailo-class NPU on matched model accuracy and input rate; report power, latency variance and integration limits.'),
    ('Reproducible validation', 'Use recorded field video plus hardware-in-the-loop PX4/SITL. Add controlled VLOS flights only after bench validation.')
], [2600, 6760])
add_para(doc, 'A defensible paper is therefore not “FPGA is faster than a GPU.” It is a systems paper about deterministic, adaptive perception for inspection, with an application-level outcome. A sensible first venue is an FPGA/embedded-systems conference or workshop; an extended version can target an embedded or UAV journal after field validation.')

heading(doc, '5. Build plan and acceptance gates', 1)
add_table(doc, ['Phase', 'Time', 'Deliverable and gate'], [
    ('0. Discovery', 'Weeks 1-2', 'Interview 5-10 inspection operators / integrators. Obtain 3+ representative recorded missions and agree a component taxonomy. Stop if no buyer confirms re-flight or image-quality pain.'),
    ('1. Software baseline', 'Weeks 3-6', 'Jetson/Python reference pipeline with target and quality labels. Establish accuracy, latency, energy and usable-observation baseline.'),
    ('2. FPGA core', 'Weeks 7-14', 'Streaming camera ingest, quantized target model and quality logic on a Kria/Zynq development platform. Gate: stable real-time operation and evidence of bounded latency.'),
    ('3. Integration MVP', 'Weeks 15-20', 'MAVLink/ROS 2 advisory interface, event log, simple operator review UI. Gate: replay and SITL demonstration end to end.'),
    ('4. Pilot', 'Weeks 21-28', 'One VLOS pilot with a service provider or integrator. Gate: measured improvement in usable capture / re-flight avoidance, not a marketing demo.'),
    ('5. Product decision', 'Weeks 29-32', 'Decide whether to package a rugged retrofit companion module or pursue an OEM/NRE design-in.')
], [1600, 1100, 6660])

heading(doc, '6. Market entry and business model', 1)
add_para(doc, 'Start with industrial drone service providers and smaller custom-platform integrators that already serve utilities. They can pilot a companion module on open PX4/ArduPilot systems, control the mission workflow, and feel the cost of re-flights directly. Do not begin by selling a generic accelerator board to OEMs or by targeting vertically integrated drone manufacturers.')
add_table(doc, ['Offer', 'Customer', 'Commercial purpose'], [
    ('Pilot kit', 'Inspection service provider / integrator', 'Development board or enclosed companion module, integration support, event logs and review export. Paid pilot or co-funded validation.'),
    ('Retrofit product', 'Selected open-platform inspection fleets', 'Ruggedized companion module with supported camera and autopilot interfaces. Hardware margin plus per-aircraft software support / analytics licence.'),
    ('OEM / NRE programme', 'Companion-computer or airframe supplier', 'Custom I/O, form factor and bitstream after the core is field-proven. This is the scale path, not the first revenue path.')
], [2100, 2650, 4610])
add_para(doc, 'The first commercial proof is not a board shipment. It is a buyer who can show fewer unusable captures, fewer repeat passes, faster analyst triage, or longer mission endurance on their own inspection workflow. That evidence will make both the paper and subsequent design-in sales materially stronger.')

heading(doc, '7. Technical architecture for v1', 1)
add_bullet(doc, 'Hardware: production-oriented Zynq UltraScale+ / Kria-class SOM for the development path; camera ingress, FPGA fabric, ARM/Linux companion application, and MAVLink/ROS 2 interface.')
add_bullet(doc, 'Data path: RGB camera -> FPGA pre-processing -> INT8 target / quality pipeline -> ARM event aggregation -> mission-controller advisory and event log. Keep raw video storage outside the timing-critical path.')
add_bullet(doc, 'Safety boundary: advisory messages are rate-limited and logged. The autopilot or human operator decides whether to alter speed, position, or capture sequence.')
add_bullet(doc, 'Future extension: sensor fusion, thermal imagery, LiDAR pre-processing and embedded capture control are v2 research/product paths once the v1 loop has a validated customer need.')

heading(doc, '8. Evaluation targets (hypotheses, not promises)', 1)
add_table(doc, ['Metric', 'Minimum thesis target', 'Commercial interpretation'], [
    ('Detector / candidate quality', 'Task-specific precision/recall established against labelled inspection imagery.', 'Avoid false confidence; candidate flags remain reviewable.'),
    ('Latency behavior', 'Measure mean, p95/p99 and observed maximum at the selected input rate.', 'A predictable advisory stream matters more than peak FPS.'),
    ('Energy', 'Report watts and joules per accepted observation across all three platforms.', 'Supports aircraft SWaP and endurance discussions honestly.'),
    ('Usable-data yield', 'Demonstrate a measurable reduction in rejected or missing target images on a held-out mission set.', 'Directly maps to reduced analyst and re-flight cost.'),
    ('Integration effort', 'Document interfaces, model update path and custom I/O changes.', 'Tests the FPGA reconfigurability argument against real engineering effort.')
], [2000, 4100, 3260])

heading(doc, '9. Key risks and how this scope controls them', 1)
add_bullet(doc, 'Dataset access risk: secure recordings and labelling agreements before hardware work. Public detection benchmarks alone do not validate capture-quality value.')
add_bullet(doc, 'Model fit risk: keep the anomaly result a candidate flag. The MVP can deliver value from component localization and usable-capture verification even if defect classification remains immature.')
add_bullet(doc, 'Hardware risk: prove the pipeline on recorded streams before live flight. Use a development SOM for the thesis; do not build a custom PCB until the pilot has validated the workflow.')
add_bullet(doc, 'Market risk: require a pilot partner to quantify a baseline re-flight, analyst-time or capture-rejection cost. If that cost is not material, pivot the same co-processor to wind-turbine or confined-space inspection with the same adaptive-capture pattern.')

heading(doc, '10. Decision summary', 1)
add_para(doc, 'The original selected vertical is good enough to keep. The change is the product and research framing: build an adaptive inspection co-processor for utility power lines, not a generic FPGA defect detector. This creates a bounded master’s-thesis scope, a credible publishable contribution, a safe MVP path, and a market entry route that does not require winning a head-to-head inference benchmark against NPUs.')

heading(doc, 'Selected evidence to verify before pilot', 2)
for source in [
    'AMD, Kria K26 SOM product page and product brief (production-grade Zynq UltraScale+ platform and interfaces): https://www.amd.com/en/products/system-on-modules/kria/k26.html',
    'EASA, Rules & Standards for drones (EU operating categories and applicable framework): https://www.easa.europa.eu/en/domains/drones-air-mobility/rules-standards',
    'Suh et al., Algorithm-Hardware Co-Optimization for Energy-Efficient Drone Detection on Resource-Constrained FPGA (baseline co-design reference): https://par.nsf.gov/servlets/purl/10322668',
    'Autonomous Inspection of Power Line Insulators with UAV on an Unmapped Transmission Tower (2026, task relevance): https://arxiv.org/abs/2602.24011',
    'AERIAL-CORE: AI-Powered Aerial Robots for Inspection and Maintenance of Electrical Power Infrastructures: https://arxiv.org/abs/2401.02343'
]: add_bullet(doc, source)

doc.core_properties.title = 'FPGA Co-Processor for Adaptive Utility-Inspection Drones'
doc.core_properties.subject = 'Revised MVP, thesis contribution and market entry plan'
doc.core_properties.author = 'CoreLink'
doc.save(OUT)
print(OUT)
